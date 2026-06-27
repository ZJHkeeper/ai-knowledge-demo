"""Ingest supported documents from ./data into a persistent Chroma collection."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_DATA_DIR = PROJECT_ROOT / "data"
DEFAULT_PERSIST_DIR = PROJECT_ROOT / "chroma_db"
DEFAULT_COLLECTION = "ai_knowledge_demo"
DEFAULT_CHUNK_SIZE = 800
DEFAULT_CHUNK_OVERLAP = 100
SUPPORTED_DOCUMENT_SUFFIXES = frozenset({".md", ".txt", ".pdf", ".docx"})
TEXT_DOCUMENT_SUFFIXES = frozenset({".md", ".txt"})
HEADING_RE = re.compile(r"^[ \t]*#{1,6}\s+.+$")
THEMATIC_BREAK_RE = re.compile(r"^[ \t]*(?:-{3,}|\*{3,}|_{3,}|(?:-[ \t]*){3,}|(?:\*[ \t]*){3,}|(?:_[ \t]*){3,})[ \t]*$")
PDF_PAGE_LABEL_RE = re.compile(r"^[ \t]*(?:page[ \t]+\d+|第[ \t]*\d+[ \t]*页)[ \t]*$", re.IGNORECASE)
CJK_SECTION_HEADING_RE = re.compile(r"^[ \t]*[一二三四五六七八九十百千]+[、.．][ \t]*\S+")
NUMBERED_SECTION_HEADING_RE = re.compile(r"^[ \t]*\d+(?:\.\d+)*[.)、][ \t]*\S+")


@dataclass(frozen=True)
class Chunk:
    """A text span ready to be stored in Chroma."""

    text: str
    metadata: dict[str, int | str]


def read_text_document_file(path: Path) -> str:
    """Read plain text documents with common UTF encodings and GB18030 fallback."""

    errors: list[str] = []
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            errors.append(f"{encoding}: {exc}")

    joined_errors = "; ".join(errors)
    raise UnicodeDecodeError(
        "document",
        b"\x00",
        0,
        1,
        f"failed to decode {path} with utf-8-sig, utf-8, or gb18030 ({joined_errors})",
    )


def read_markdown_file(path: Path) -> str:
    """Read Markdown text with common UTF encodings and GB18030 fallback."""

    return read_text_document_file(path)


def read_document_file(path: Path) -> str:
    """Read a supported document file as plain text."""

    suffix = path.suffix.lower()
    if suffix == ".md":
        return read_text_document_file(path)
    if suffix == ".txt":
        return _normalize_extracted_text(read_text_document_file(path))
    if suffix == ".pdf":
        return _read_pdf_file(path)
    if suffix == ".docx":
        return _normalize_extracted_text(_read_docx_file(path))
    raise ValueError(f"unsupported document type: {path}")


def discover_document_files(data_dir: Path) -> list[Path]:
    """Find supported document files in a deterministic order."""

    if not data_dir.exists():
        return []
    return sorted(
        (
            path
            for path in data_dir.rglob("*")
            if path.is_file() and path.suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES
        ),
        key=lambda path: path.as_posix().lower(),
    )


def discover_markdown_files(data_dir: Path) -> list[Path]:
    """Find supported document files in a deterministic order."""

    return discover_document_files(data_dir)


def chunk_markdown(
    text: str,
    source: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[Chunk]:
    """Split Markdown into Chroma-ready chunks with source metadata."""

    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap must be greater than or equal to 0")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size")

    chunks: list[Chunk] = []
    for start, end in _markdown_spans(text):
        chunks.extend(_split_span(text, start, end, source, chunk_size, chunk_overlap))

    return [
        Chunk(
            text=chunk.text,
            metadata={**chunk.metadata, "chunk_index": index},
        )
        for index, chunk in enumerate(chunks)
    ]


def build_chunks_for_files(
    document_files: Iterable[Path],
    data_dir: Path,
    chunk_size: int,
    chunk_overlap: int,
) -> tuple[list[Chunk], int, list[str]]:
    """Read files and build chunks, returning skipped-file error messages too."""

    all_chunks: list[Chunk] = []
    files_read = 0
    errors: list[str] = []

    for document_file in document_files:
        source = document_file.relative_to(data_dir).as_posix()
        try:
            text = read_document_file(document_file)
        except Exception as exc:
            errors.append(f"{document_file}: {exc}")
            continue

        file_chunks = chunk_markdown(text, source, chunk_size, chunk_overlap)
        all_chunks.extend(file_chunks)
        files_read += 1

    return all_chunks, files_read, errors


def ingest_chunks(chunks: list[Chunk], persist_dir: Path, collection_name: str) -> None:
    """Write chunks into Chroma, replacing chunks from the same source."""

    import chromadb

    persist_dir.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(persist_dir))
    collection = client.get_or_create_collection(name=collection_name)

    sources = sorted({str(chunk.metadata["source"]) for chunk in chunks})
    for source in sources:
        collection.delete(where={"source": source})

    if not chunks:
        return

    collection.add(
        ids=[_chunk_id(chunk) for chunk in chunks],
        documents=[chunk.text for chunk in chunks],
        metadatas=[chunk.metadata for chunk in chunks],
    )


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments."""

    parser = argparse.ArgumentParser(
        description="Ingest supported documents from ./data into a Chroma vector database."
    )
    parser.add_argument("--data-dir", type=Path, default=DEFAULT_DATA_DIR)
    parser.add_argument("--persist-dir", type=Path, default=DEFAULT_PERSIST_DIR)
    parser.add_argument("--collection", default=DEFAULT_COLLECTION)
    parser.add_argument("--chunk-size", type=int, default=DEFAULT_CHUNK_SIZE)
    parser.add_argument("--chunk-overlap", type=int, default=DEFAULT_CHUNK_OVERLAP)
    return parser.parse_args()


def main() -> int:
    """Run document ingestion."""

    args = parse_args()
    data_dir = args.data_dir.resolve()
    persist_dir = args.persist_dir.resolve()

    document_files = discover_document_files(data_dir)
    chunks, files_read, errors = build_chunks_for_files(
        document_files,
        data_dir,
        args.chunk_size,
        args.chunk_overlap,
    )
    ingest_chunks(chunks, persist_dir, args.collection)

    for error in errors:
        print(f"Skipped file: {error}")

    print(f"Files read: {files_read}")
    print(f"Chunks written: {len(chunks)}")
    print(f"Collection: {args.collection}")
    print(f"Chroma path: {persist_dir}")
    return 0


def _markdown_spans(text: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    span_start = 0

    for line in re.finditer(r".*(?:\r?\n|$)", text):
        if line.start() == line.end():
            continue

        line_text = line.group(0).rstrip("\r\n")
        if HEADING_RE.match(line_text):
            trimmed = _trim_span(text, span_start, line.start())
            if trimmed is not None:
                spans.append(trimmed)
            span_start = line.start()
        elif THEMATIC_BREAK_RE.match(line_text):
            trimmed = _trim_span(text, span_start, line.start())
            if trimmed is not None:
                spans.append(trimmed)
            span_start = line.end()

    trimmed = _trim_span(text, span_start, len(text))
    if trimmed is not None:
        spans.append(trimmed)
    return spans


def _read_pdf_file(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    page_texts = []
    for page in reader.pages:
        page_text = _normalize_extracted_text(page.extract_text() or "")
        if page_text:
            page_texts.append(page_text)
    return "\n\n---\n\n".join(page_texts).strip()


def _normalize_extracted_text(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or PDF_PAGE_LABEL_RE.match(line):
            continue
        if _looks_like_extracted_heading(line) and not HEADING_RE.match(line):
            line = f"## {line}"
        lines.append(line)
    return "\n".join(lines).strip()


def _looks_like_extracted_heading(line: str) -> bool:
    return bool(
        CJK_SECTION_HEADING_RE.match(line)
        or NUMBERED_SECTION_HEADING_RE.match(line)
    )


def _read_docx_file(path: Path) -> str:
    from docx import Document

    document = Document(path)
    blocks: list[str] = []
    blocks.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                if cell_text:
                    cells.append(cell_text)
            if cells:
                blocks.append("\t".join(cells))

    return "\n\n".join(blocks).strip()


def _split_span(
    text: str,
    start: int,
    end: int,
    source: str,
    chunk_size: int,
    chunk_overlap: int,
) -> list[Chunk]:
    chunks: list[Chunk] = []
    position = start

    while position < end:
        chunk_end = min(position + chunk_size, end)
        if chunk_end < end:
            chunk_end = _best_break(text, position, chunk_end)

        trimmed = _trim_span(text, position, chunk_end)
        if trimmed is not None:
            trimmed_start, trimmed_end = trimmed
            chunks.append(
                Chunk(
                    text=text[trimmed_start:trimmed_end],
                    metadata={
                        "source": source,
                        "chunk_index": 0,
                        "start_char": trimmed_start,
                        "end_char": trimmed_end,
                    },
                )
            )

        if chunk_end >= end:
            break

        next_position = max(start, chunk_end - chunk_overlap)
        if next_position <= position:
            next_position = chunk_end
        position = next_position

    return chunks


def _best_break(text: str, start: int, default_end: int) -> int:
    minimum = start + max(1, int((default_end - start) * 0.75))
    for marker in ("\n\n", "\n", " "):
        index = text.rfind(marker, minimum, default_end)
        if index != -1:
            return index + len(marker)
    return default_end


def _trim_span(text: str, start: int, end: int) -> tuple[int, int] | None:
    while start < end and text[start].isspace():
        start += 1
    while end > start and text[end - 1].isspace():
        end -= 1
    if start >= end:
        return None
    return start, end


def _chunk_id(chunk: Chunk) -> str:
    source = str(chunk.metadata["source"]).replace("\\", "/")
    index = int(chunk.metadata["chunk_index"])
    return f"{source}:{index}"


if __name__ == "__main__":
    raise SystemExit(main())
