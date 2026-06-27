import json
import sys
import tempfile
import unittest
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = PROJECT_ROOT / "src"
sys.path.insert(0, str(SRC_DIR))

from ai_knowledge_demo.ingest import (
    Chunk,
    build_bm25_index,
    chunk_markdown,
    default_bm25_index_path,
    discover_document_files,
    read_document_file,
    tokenize_for_bm25,
    write_bm25_index,
)


def write_simple_pdf(path: Path, text: str) -> None:
    lines = text.splitlines() or [text]
    escaped_lines = [
        line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        for line in lines
    ]
    text_ops = " T* ".join(f"({line}) Tj" for line in escaped_lines)
    content = f"BT /F1 12 Tf 72 720 Td 14 TL {text_ops} ET"
    content_bytes = content.encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content_bytes)).encode("ascii") + b" >>\nstream\n"
        + content_bytes
        + b"\nendstream",
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        f"trailer << /Root 1 0 R /Size {len(objects) + 1} >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(pdf)


class IngestChunkTests(unittest.TestCase):
    def test_discovers_supported_documents_in_stable_order(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            (data_dir / "b.TXT").write_text("text", encoding="utf-8")
            (data_dir / "a.md").write_text("markdown", encoding="utf-8")
            (data_dir / "nested").mkdir()
            (data_dir / "nested" / "c.pdf").write_bytes(b"%PDF-1.4\n")
            (data_dir / "nested" / "d.docx").write_bytes(b"placeholder")
            (data_dir / "ignored.csv").write_text("nope", encoding="utf-8")

            discovered = discover_document_files(data_dir)

        self.assertEqual(
            [path.relative_to(data_dir).as_posix() for path in discovered],
            ["a.md", "b.TXT", "nested/c.pdf", "nested/d.docx"],
        )

    def test_reads_markdown_and_txt_with_encoding_fallbacks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            data_dir = Path(temp_dir)
            markdown_file = data_dir / "policy.md"
            text_file = data_dir / "rules.txt"
            markdown_file.write_text("# Title\n\nUTF-8 text", encoding="utf-8-sig")
            text_file.write_text("中文文本", encoding="gb18030")

            self.assertEqual(read_document_file(markdown_file), "# Title\n\nUTF-8 text")
            self.assertEqual(read_document_file(text_file), "中文文本")

    def test_txt_numbered_sections_become_separate_chunks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            text_file = Path(temp_dir) / "rules.txt"
            text_file.write_text(
                "\n".join(
                    [
                        "物流异常与签收争议处理规则",
                        "",
                        "一、物流状态异常",
                        "超过 7 天没有更新可申请催单。",
                        "",
                        "二、签收争议处理结果",
                        "核实运输丢失后可补发或退款。",
                    ]
                ),
                encoding="utf-8",
            )

            text = read_document_file(text_file)
            chunks = chunk_markdown(text, text_file.name)

        self.assertEqual(
            [chunk.text for chunk in chunks],
            [
                "物流异常与签收争议处理规则",
                "## 一、物流状态异常\n超过 7 天没有更新可申请催单。",
                "## 二、签收争议处理结果\n核实运输丢失后可补发或退款。",
            ],
        )

    def test_reads_docx_paragraphs_and_tables(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            docx_file = Path(temp_dir) / "manual.docx"
            document = Document()
            document.add_paragraph("Membership invoice paragraph.")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Invoice type"
            table.cell(0, 1).text = "Electronic"
            document.save(docx_file)

            text = read_document_file(docx_file)

        self.assertIn("Membership invoice paragraph.", text)
        self.assertIn("Invoice type\tElectronic", text)

    def test_docx_numbered_sections_become_separate_chunks(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            docx_file = Path(temp_dir) / "manual.docx"
            document = Document()
            document.add_paragraph("会员与发票处理补充规则")
            document.add_paragraph("一、会员自动续费")
            document.add_paragraph("自动续费后 24 小时内未使用权益可申请退款。")
            document.add_paragraph("二、电子发票")
            document.add_paragraph("订单退款后，对应电子发票会自动作废。")
            document.save(docx_file)

            text = read_document_file(docx_file)
            chunks = chunk_markdown(text, docx_file.name)

        self.assertEqual(
            [chunk.text for chunk in chunks],
            [
                "会员与发票处理补充规则",
                "## 一、会员自动续费\n自动续费后 24 小时内未使用权益可申请退款。",
                "## 二、电子发票\n订单退款后，对应电子发票会自动作废。",
            ],
        )

    def test_reads_pdf_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_file = Path(temp_dir) / "manual.pdf"
            write_simple_pdf(pdf_file, "International payment refund manual.")

            text = read_document_file(pdf_file)

        self.assertIn("International payment refund manual.", text)

    def test_pdf_numbered_sections_become_separate_chunks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_file = Path(temp_dir) / "manual.pdf"
            write_simple_pdf(
                pdf_file,
                "\n".join(
                    [
                        "page 1",
                        "International payment refund manual",
                        "1. Refund timing",
                        "Refunds usually arrive in 7-15 business days.",
                        "2. Manual review",
                        "High-risk international orders need manual review.",
                    ]
                ),
            )

            text = read_document_file(pdf_file)
            chunks = chunk_markdown(text, "manual.pdf")

        self.assertNotIn("page 1", text)
        self.assertIn("## 1. Refund timing", text)
        self.assertIn("## 2. Manual review", text)
        self.assertEqual(
            [chunk.text for chunk in chunks],
            [
                "International payment refund manual",
                "## 1. Refund timing\nRefunds usually arrive in 7-15 business days.",
                "## 2. Manual review\nHigh-risk international orders need manual review.",
            ],
        )

    def test_short_markdown_stays_in_one_chunk(self) -> None:
        text = "# Refunds\n\nRefund requests must be filed within 7 days."

        chunks = chunk_markdown(text, "refund_policy.md", chunk_size=800, chunk_overlap=100)

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].text, text)
        self.assertEqual(chunks[0].metadata["source"], "refund_policy.md")
        self.assertEqual(chunks[0].metadata["chunk_index"], 0)
        self.assertEqual(chunks[0].metadata["start_char"], 0)
        self.assertEqual(chunks[0].metadata["end_char"], len(text))

    def test_long_paragraph_uses_size_and_overlap(self) -> None:
        text = "a" * 25

        chunks = chunk_markdown(text, "long.md", chunk_size=10, chunk_overlap=3)

        self.assertEqual([chunk.text for chunk in chunks], ["a" * 10, "a" * 10, "a" * 10, "a" * 4])
        self.assertEqual(
            [(chunk.metadata["start_char"], chunk.metadata["end_char"]) for chunk in chunks],
            [(0, 10), (7, 17), (14, 24), (21, 25)],
        )

    def test_metadata_is_stable_across_markdown_sections(self) -> None:
        text = "# One\n\nFirst section.\n\n---\n\n## Two\n\nSecond section."

        chunks = chunk_markdown(text, "nested/refund_policy.md", chunk_size=800, chunk_overlap=100)

        self.assertEqual([chunk.metadata["chunk_index"] for chunk in chunks], [0, 1])
        self.assertEqual(
            [chunk.metadata["source"] for chunk in chunks],
            ["nested/refund_policy.md", "nested/refund_policy.md"],
        )
        self.assertEqual(chunks[0].text, "# One\n\nFirst section.")
        self.assertEqual(chunks[1].text, "## Two\n\nSecond section.")

    def test_thematic_breaks_only_split_sections(self) -> None:
        text = "# One\n\nFirst\n\n***\n\n## Two\n\nSecond\n\n___\n\n## Three\n\nThird\n\n- - -\n\n## Four\n\nFourth"

        chunks = chunk_markdown(text, "breaks.md", chunk_size=800, chunk_overlap=100)

        self.assertEqual(
            [chunk.text for chunk in chunks],
            [
                "# One\n\nFirst",
                "## Two\n\nSecond",
                "## Three\n\nThird",
                "## Four\n\nFourth",
            ],
        )
        self.assertFalse(any(chunk.text.strip() in {"---", "***", "___", "- - -"} for chunk in chunks))

    def test_bm25_index_preserves_chunks_metadata_and_tokens(self) -> None:
        chunks = [
            Chunk(
                text="Visa 信用卡退款通常需要 7-15 个工作日。",
                metadata={
                    "source": "refund_policy.md",
                    "chunk_index": 2,
                    "start_char": 10,
                    "end_char": 42,
                },
            )
        ]

        index = build_bm25_index(chunks)

        self.assertEqual(index["version"], 1)
        self.assertEqual(len(index["chunks"]), 1)
        self.assertEqual(index["chunks"][0]["id"], "refund_policy.md:2")
        self.assertEqual(index["chunks"][0]["metadata"]["source"], "refund_policy.md")
        self.assertEqual(index["chunks"][0]["metadata"]["chunk_index"], 2)
        self.assertIn("visa", index["chunks"][0]["tokens"])
        self.assertIn("退款", index["chunks"][0]["tokens"])
        self.assertIn("信用卡", index["chunks"][0]["tokens"])

    def test_write_bm25_index_uses_default_path_and_writes_json(self) -> None:
        chunks = [
            Chunk(
                text="Refunds arrive in 3-5 business days.",
                metadata={"source": "refund_policy.md", "chunk_index": 0},
            )
        ]

        with tempfile.TemporaryDirectory() as temp_dir:
            index_path = default_bm25_index_path(Path(temp_dir))
            write_bm25_index(chunks, index_path)
            data = json.loads(index_path.read_text(encoding="utf-8"))

        self.assertEqual(index_path.name, "bm25_index.json")
        self.assertEqual(len(data["chunks"]), 1)
        self.assertEqual(data["chunks"][0]["metadata"]["chunk_index"], 0)

    def test_bm25_tokenizer_expands_card_brands(self) -> None:
        tokens = tokenize_for_bm25("Visa 退款需要多久？")

        self.assertIn("visa", tokens)
        self.assertIn("信用卡", tokens)
        self.assertIn("国际银行卡", tokens)
        self.assertIn("退款", tokens)


if __name__ == "__main__":
    unittest.main()
